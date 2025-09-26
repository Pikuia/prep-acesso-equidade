# -*- coding: utf-8 -*-
"""
PrEP — Entre a Oferta e o Conhecimento (Município de São Paulo)

App Streamlit em arquivo único, com:
- Banco SQLite (coleta via formulário)
- Integração ao Painel PrEP (download .zip) [MS]
- Modelos: Regressão Logística (OR/IC95%) e Perfis Latentes (GaussianMixture+BIC)
- Métrica de GAP (share PrEP - share HIV) com IC95% por diferença de proporções
- Mapas por Subprefeituras (geopandas) + opção de enviar seu GeoJSON/SHP oficial
- Exportação de gráficos (PNG) e Relatório .docx com as figuras
"""

from __future__ import annotations
import os, io, zipfile, unicodedata, warnings, base64
from pathlib import Path
from datetime import datetime
import requests
import sqlite3
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st

# viz
import plotly.express as px

# modelos
import statsmodels.api as sm
from sklearn.preprocessing import OneHotEncoder
from sklearn.mixture import GaussianMixture

# geoespacial
import geopandas as gpd

# relatório
from docx import Document
from docx.shared import Inches

warnings.filterwarnings("ignore", category=FutureWarning)

# ======================= CONFIGURAÇÕES =======================
st.set_page_config(page_title="PrEP — Entre a Oferta e o Conhecimento", page_icon="🧭", layout="wide")
BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
EXPORT_DIR = DATA_DIR / "exports"
for d in (DATA_DIR, EXPORT_DIR):
    d.mkdir(exist_ok=True)

SQLITE_PATH = DATA_DIR / "prep_research.db"

# URLs oficiais
URL_PREP_ZIP = "https://mediacenter.aids.gov.br/prep/Dados_PrEP_transparencia.zip"  # Painel PrEP (MS) — base pública
URL_SUBPREF_GEOJSON = "https://raw.githubusercontent.com/codigurbano/subprefeituras-sp/master/data/subprefeituras-sp.geojson"
GEOJSON_LOCAL = DATA_DIR / "geodata_subprefeituras.geojson"

# ======================= FUNÇÕES CORE =======================
def strip_accents(s: str) -> str:
    if not isinstance(s, str):
        return s
    return ''.join(c for c in unicodedata.normalize('NFKD', s) if not unicodedata.combining(c))

def ensure_sqlite():
    DDL = """
    PRAGMA foreign_keys = ON;

    CREATE TABLE IF NOT EXISTS respostas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ts_utc TEXT NOT NULL,
        versao_form TEXT NOT NULL,
        conhecimento_prep TEXT,
        conhecimento_pep TEXT,
        acesso_servicos TEXT,
        fonte_informacao TEXT,
        uso_preppep TEXT,
        conhece_usuarios TEXT,
        teste_hiv_freq TEXT,
        metodos_prevencao TEXT,
        genero TEXT,
        orientacao_sexual TEXT,
        raca TEXT,
        faixa_etaria TEXT,
        escolaridade TEXT,
        renda TEXT,
        regiao TEXT,
        subprefeitura TEXT,
        comentarios TEXT
    );

    CREATE TABLE IF NOT EXISTS execucoes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ts_utc TEXT NOT NULL,
        acao TEXT NOT NULL,
        detalhes TEXT
    );
    """
    conn = sqlite3.connect(str(SQLITE_PATH), check_same_thread=False)
    cur = conn.cursor()
    for stmt in DDL.strip().split(";\n\n"):
        if stmt.strip():
            cur.execute(stmt)
    conn.commit()
    conn.close()

def insert_resposta(payload: dict, versao_form: str = "v2"):
    conn = sqlite3.connect(str(SQLITE_PATH))
    payload_db = {
        'ts_utc': datetime.utcnow().isoformat(timespec='seconds'),
        'versao_form': versao_form,
        'conhecimento_prep': payload.get('Conhecimento_PrEP'),
        'conhecimento_pep': payload.get('Conhecimento_PEP'),
        'acesso_servicos': payload.get('Acesso_servicos'),
        'fonte_informacao': payload.get('Fonte_informacao'),
        'uso_preppep': payload.get('Uso_PrepPEP'),
        'conhece_usuarios': payload.get('Conhece_usuarios'),
        'teste_hiv_freq': payload.get('Teste_HIV_frequencia'),
        'metodos_prevencao': payload.get('Metodos_prevencao'),
        'genero': payload.get('Genero'),
        'orientacao_sexual': payload.get('Orientacao_sexual'),
        'raca': payload.get('Raca'),
        'faixa_etaria': payload.get('Faixa_etaria'),
        'escolaridade': payload.get('Escolaridade'),
        'renda': payload.get('Renda'),
        'regiao': payload.get('Regiao'),
        'subprefeitura': payload.get('Subprefeitura'),
        'comentarios': payload.get('Comentarios')
    }
    cols = ",".join(payload_db.keys())
    qmarks = ",".join(["?"] * len(payload_db))
    conn.execute(f"INSERT INTO respostas ({cols}) VALUES ({qmarks})", list(payload_db.values()))
    conn.commit()
    conn.close()

def load_respostas() -> pd.DataFrame:
    conn = sqlite3.connect(str(SQLITE_PATH))
    df = pd.read_sql_query("SELECT * FROM respostas ORDER BY id DESC", conn)
    conn.close()
    return df

@st.cache_data(show_spinner=False)
def download_painel_prep() -> tuple[pd.DataFrame, pd.DataFrame]:
    r = requests.get(URL_PREP_ZIP, timeout=120)
    r.raise_for_status()
    z = zipfile.ZipFile(io.BytesIO(r.content))
    names = [n for n in z.namelist() if n.lower().endswith(".csv")]
    fn_usuarios = next((n for n in names if "usuario" in n.lower()), None)
    fn_dispensa = next((n for n in names if "disp" in n.lower()), None)
    if not fn_usuarios:
        raise ValueError(f"CSV de usuários não encontrado. Arquivos: {names[:5]} ...")

    df_usuarios = pd.read_csv(z.open(fn_usuarios), sep=";", encoding="utf-8")
    df_usuarios.columns = [c.strip().upper() for c in df_usuarios.columns]
    if fn_dispensa:
        df_disp = pd.read_csv(z.open(fn_dispensa), sep=";", encoding="utf-8")
        df_disp.columns = [c.strip().upper() for c in df_disp.columns]
    else:
        df_disp = pd.DataFrame()
    return df_usuarios, df_disp

def filter_municipio_sp(df_usuarios: pd.DataFrame) -> pd.DataFrame:
    return df_usuarios[
        (df_usuarios.get('UF') == 'SP') &
        (df_usuarios.get('MUNICIPIO', '').str.upper().str.contains('SÃO PAULO', na=False))
    ].copy()

def harmonize_survey(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df['Conhecimento_PrEP_bin'] = df['conhecimento_prep'].map({
        "Sim, conheço bem": 1, "Conheço parcialmente": 1,
        "Já ouvi falar mas não sei detalhes": 0, "Não conheço": 0
    }).astype('Int64')

    uso_map = {"Sim, uso atualmente": 1}
    df['Uso_PrEP_bin'] = df['uso_preppep'].map(uso_map).fillna(0).astype(int)
    df.loc[df['metodos_prevencao'].fillna('').str.contains('PrEP', case=False), 'Uso_PrEP_bin'] = 1

    for c in ['raca', 'faixa_etaria', 'escolaridade', 'genero', 'orientacao_sexual', 'regiao', 'subprefeitura']:
        if c in df.columns:
            df[c] = df[c].astype(str).str.strip()
    return df

def logistic_or(df: pd.DataFrame, target: str, features: list[str]) -> tuple[sm.Logit, pd.DataFrame]:
    data = df[[target] + features].dropna()
    if data.empty:
        raise ValueError("Dados insuficientes para a regressão.")
    X = pd.get_dummies(data[features], drop_first=True)
    X = sm.add_constant(X)
    y = data[target].astype(int)
    model = sm.Logit(y, X).fit(disp=0)
    params = model.params
    conf = model.conf_int()
    or_df = pd.DataFrame({
        'feature': params.index,
        'OR': np.exp(params.values),
        'CI_low': np.exp(conf[0].values),
        'CI_high': np.exp(conf[1].values),
        'pvalue': model.pvalues.values
    })
    return model, or_df.sort_values('OR', ascending=False)

def latent_profiles(df: pd.DataFrame, cat_cols: list[str], k_range=range(2, 6)) -> tuple[pd.DataFrame, dict]:
    enc = OneHotEncoder(sparse_output=False, handle_unknown='ignore')
    Z = enc.fit_transform(df[cat_cols].astype(str))
    best = {'bic': float('inf'), 'k': None, 'model': None}
    for k in k_range:
        gm = GaussianMixture(n_components=k, covariance_type='spherical', random_state=42)
        gm.fit(Z)
        bic = gm.bic(Z)
        if bic < best['bic']:
            best = {'bic': bic, 'k': k, 'model': gm}
    labels = best['model'].predict(Z)
    out = df.copy()
    out['Perfil_Latente'] = labels
    return out, best

def gap_repr_with_ci(df_prep_sp: pd.DataFrame, df_cases_agg: pd.DataFrame, dim_col: str, ano: int) -> pd.DataFrame:
    """
    Calcula GAP = share_prep - share_hiv e IC95% para a diferença de proporções (aprox. normal)
    p1 = x1/n1 (PrEP), p2 = y1/n2 (HIV); IC95%: (p1-p2) ± 1.96*sqrt(p1*(1-p1)/n1 + p2*(1-p2)/n2)
    """
    a = df_prep_sp[df_prep_sp['ANO'] == ano].groupby(dim_col).size().rename('prep').reset_index()
    b = df_cases_agg[df_cases_agg['ANO'] == ano].groupby(dim_col)['COUNT'].sum().rename('hiv').reset_index()
    m = a.merge(b, on=dim_col, how='outer').fillna(0)

    n1 = m['prep'].sum()
    n2 = m['hiv'].sum()
    m['p1'] = m['prep'] / n1 if n1 > 0 else 0.0
    m['p2'] = m['hiv'] / n2 if n2 > 0 else 0.0
    m['gap_repr'] = m['p1'] - m['p2']

    # erro padrão e IC
    se = np.sqrt((m['p1'] * (1 - m['p1']) / max(n1, 1)) + (m['p2'] * (1 - m['p2']) / max(n2, 1)))
    z = 1.96
    m['ci_low'] = m['gap_repr'] - z * se
    m['ci_high'] = m['gap_repr'] + z * se
    m['signif'] = np.where((m['ci_low'] > 0) | (m['ci_high'] < 0), "Significativo (95%)", "Não significativo")
    return m.sort_values('gap_repr')

@st.cache_data(show_spinner=False)
def ensure_geodata(default=True, uploaded_file=None) -> gpd.GeoDataFrame:
    """
    Retorna GeoDataFrame de Subprefeituras.
    - default=True: usa espelho público (rápido).
    - uploaded_file: permite substituir por GeoJSON/SHP oficial do GeoSampa.
    """
    if uploaded_file is not None:
        # Detecta extensão; aceita GeoJSON, SHP (num .zip), GPKG (se desejar)
        name = uploaded_file.name.lower()
        if name.endswith('.geojson') or name.endswith('.json'):
            gdf = gpd.read_file(uploaded_file)
        elif name.endswith('.zip'):
            # espera um shapefile zipado
            import tempfile, zipfile as zf
            tmpdir = tempfile.mkdtemp()
            with zf.ZipFile(uploaded_file) as z:
                z.extractall(tmpdir)
            gdf = gpd.read_file(tmpdir)
        else:
            gdf = gpd.read_file(uploaded_file)
    else:
        if not GEOJSON_LOCAL.exists():
            r = requests.get(URL_SUBPREF_GEOJSON, timeout=90)
            r.raise_for_status()
            GEOJSON_LOCAL.write_bytes(r.content)
        gdf = gpd.read_file(GEOJSON_LOCAL)

    # Detecta coluna de nome
    name_col = None
    for col in gdf.columns:
        if col.lower() in ['subpref', 'subprefeitura', 'sp', 'nome', 'nm_subpref']:
            name_col = col
            break
    if name_col is None:
        name_col = gdf.columns[0]
    gdf = gdf.rename(columns={name_col: 'SUBPREF'})
    gdf['SUBPREF'] = gdf['SUBPREF'].astype(str)
    return gdf

def choropleth_subpref(gdf_subpref: gpd.GeoDataFrame, df_metric: pd.DataFrame,
                       key_df='Subprefeitura', key_gdf='SUBPREF', value_col='valor') -> gpd.GeoDataFrame:
    g = gdf_subpref.copy()
    left = df_metric.copy()
    left[key_df] = left[key_df].astype(str).str.upper().apply(strip_accents)
    g[key_gdf] = g[key_gdf].astype(str).str.upper().apply(strip_accents)
    g = g.merge(left[[key_df, value_col]], left_on=key_gdf, right_on=key_df, how='left')
    return g

def save_fig(fig, filename: str) -> str:
    path = EXPORT_DIR / filename
    # plotly fig
    if hasattr(fig, "to_image"):
        fig.write_image(str(path))
    else:
        # matplotlib fig
        fig.savefig(str(path), dpi=150, bbox_inches='tight')
    return str(path)

def build_report(path_docx: str, resumo_bullets: list[str], figuras_png: list[str]) -> None:
    doc = Document()
    doc.add_heading('PrEP – Entre a Oferta e o Conhecimento (Município de São Paulo)', level=1)
    doc.add_paragraph('Resumo executivo')
    for b in resumo_bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_heading('Gráficos', level=2)
    for fig in figuras_png:
        p = doc.add_paragraph()
        p.add_run().add_picture(fig, width=Inches(6.5))
    doc.add_heading('Fontes de dados', level=2)
    doc.add_paragraph('Painel PrEP (MS): https://www.gov.br/aids/pt-br/indicadores-epidemiologicos/painel-de-monitoramento/painel-prep')
    doc.add_paragraph('Painéis de Indicadores e Dados Básicos (MS): https://www.gov.br/aids/pt-br/indicadores-epidemiologicos/paineis-de-indicadores-e-dados-basicos')
    doc.add_paragraph('Boletim Epidemiológico HIV e Aids (MS)')
    doc.add_paragraph('PrEP – Prefeitura de São Paulo (SPrEP/unidades)')
    doc.save(path_docx)

def download_button_from_file(filepath: str, label: str):
    with open(filepath, "rb") as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:file/octet-stream;base64,{b64}" download="{Path(filepath).name}">{label}</a>'
    st.markdown(href, unsafe_allow_html=True)

# ======================= INTERFACE =======================
ensure_sqlite()
st.title("PrEP — Entre a Oferta e o Conhecimento (São Paulo)")

menu = st.sidebar.radio(
    "Navegação",
    ["Coleta", "Integração de Dados", "Modelos", "Gap por Grupo", "Mapas", "Relatório", "Sobre"]
)

# ----------------- COLETA -----------------
if menu == "Coleta":
    st.header("Questionário — Conhecimento/uso de PrEP/PEP")
    with st.form("frm"):
        col1, col2 = st.columns(2)
        with col1:
            q1 = st.radio("Você conhece a PrEP?",
                          ["Sim, conheço bem", "Conheço parcialmente", "Já ouvi falar mas não sei detalhes", "Não conheço"])
            q2 = st.radio("E a PEP?",
                          ["Sim, conheço bem", "Conheço parcialmente", "Já ouvi falar mas não sei detalhes", "Não conheço"])
            q3 = st.radio("Você sabe onde conseguir PrEP/PEP em São Paulo?",
                          ["Sim, conheço vários serviços", "Conheço apenas um local", "Não sei mas gostaria de saber", "Não sei e não tenho interesse"])
            q4 = st.selectbox("Como soube da PrEP/PEP?",
                              ["Profissional de saúde", "Amigos/conhecidos", "Internet/redes sociais", "Material informativo", "Nunca ouvi falar", "Outra"])
            q5 = st.radio("Você já usou/usa PrEP/PEP?",
                          ["Sim, uso atualmente", "Sim, já usei no passado", "Não, mas pretendo usar", "Não uso e não tenho interesse"])
            q6 = st.radio("Conhece alguém que usa PrEP/PEP?",
                          ["Sim, vários", "Sim, alguns", "Não conheço ninguém", "Prefiro não responder"])
        with col2:
            q7 = st.radio("Frequência de teste de HIV",
                          ["A cada 3 meses", "A cada 6 meses", "Uma vez por ano", "Raramente faço", "Nunca fiz", "Prefiro não responder"])
            genero = st.selectbox("Identidade de gênero",
                                  ["Mulher cisgênero", "Homem cisgênero", "Mulher trans", "Homem trans", "Pessoa não-binária", "Travesti", "Outro", "Prefiro não responder"])
            orient = st.selectbox("Orientação sexual",
                                  ["Assexual", "Bissexual", "Gay", "Lésbica", "Pansexual", "Heterossexual", "Queer", "Outra", "Prefiro não responder"])
            raca = st.radio("Raça/cor", ["Amarela", "Branca", "Indígena", "Parda", "Preta", "Prefiro não responder"])
            faixa = st.radio("Faixa etária", ["13-17", "18-24", "25-29", "30-39", "40-49", "50-59", "60+", "Prefiro não responder"])
            escolar = st.selectbox("Escolaridade",
                                   ["Fundamental incompleto", "Fundamental completo", "Médio incompleto", "Médio completo",
                                    "Superior incompleto", "Superior completo", "Pós-graduação", "Prefiro não responder"])
            renda = st.radio("Renda mensal individual",
                             ["Até 1 SM", "1-2 SM", "2-3 SM", "3-5 SM", "Mais de 5 SM", "Prefiro não responder"])
            regiao = st.selectbox("Região de São Paulo onde mora",
                                  ["Centro expandido", "Zona Norte", "Zona Sul", "Zona Leste", "Zona Oeste",
                                   "Região Metropolitana", "Não moro em S. Paulo", "Prefiro não responder"])
            subpref = st.text_input("Subprefeitura (opcional, para mapas)")

        metodos = st.multiselect("Quais métodos de prevenção você utiliza?",
                                 ["PrEP", "PEP", "Camisinha masculina", "Camisinha feminina", "Testagem regular", "Não utilizo", "Outro"])
        comentarios = st.text_area("Comentários (opcional)", height=80)
        enviado = st.form_submit_button("Enviar")

    if enviado:
        payload = {
            "Conhecimento_PrEP": q1, "Conhecimento_PEP": q2, "Acesso_servicos": q3, "Fonte_informacao": q4,
            "Uso_PrepPEP": q5, "Conhece_usuarios": q6, "Teste_HIV_frequencia": q7,
            "Metodos_prevencao": ", ".join([m for m in metodos]),
            "Genero": genero, "Orientacao_sexual": orient, "Raca": raca, "Faixa_etaria": faixa,
            "Escolaridade": escolar, "Renda": renda, "Regiao": regiao,
            "Subprefeitura": subpref.strip() if subpref else None, "Comentarios": comentarios.strip() if comentarios else None
        }
        insert_resposta(payload)
        st.success("✅ Resposta registrada. Obrigado por contribuir!")
    st.divider()
    df = load_respostas()
    st.caption(f"Total de respostas: {len(df)}")
    st.dataframe(df.head(30), use_container_width=True)

# ----------------- INTEGRAÇÃO -----------------
elif menu == "Integração de Dados":
    st.header("Integração com dados públicos (Painel PrEP, MS)")
    st.write("Baixe a base pública do Painel PrEP (2018–2024) e filtre o Município de São Paulo.")
    if st.button("Baixar e integrar a base do Painel PrEP"):
        with st.spinner("Baixando base do Painel PrEP..."):
            try:
                df_users, df_disp = download_painel_prep()
                df_sp = filter_municipio_sp(df_users)
                st.session_state['PREP_SP'] = df_sp
                st.success(f"Usuários (Município de São Paulo): {len(df_sp):,}")
                st.dataframe(df_sp.head(30), use_container_width=True)
                st.info("Fonte: Painel PrEP — Ministério da Saúde. Consulte metodologia/definições no painel.")
            except Exception as e:
                st.error(f"Erro ao baixar: {e}")

    st.markdown("### Proxy de necessidade (incidência/positividade/casos)")
    st.write("Faça upload de um CSV **agregado** com colunas: `ANO`, `DIM_COL` (ex.: `RACA_COR`), `COUNT`.")
    up = st.file_uploader("CSV de casos/positividade (proxy de necessidade)", type=['csv'])
    if up:
        df_cases = pd.read_csv(up)
        st.session_state['CASES_REF_RAW'] = df_cases
        st.dataframe(df_cases.head(15), use_container_width=True)
        st.success("CSV carregado.")

# ----------------- MODELOS -----------------
elif menu == "Modelos":
    st.header("Modelos — Regressão Logística e Perfis Latentes")
    df = harmonize_survey(load_respostas())
    if len(df) < 50:
        st.info("Precisamos de pelo menos ~50 respostas para resultados mais estáveis.")
    else:
        features = st.multiselect(
            "Variáveis preditoras",
            ['genero', 'orientacao_sexual', 'raca', 'faixa_etaria', 'escolaridade', 'renda', 'regiao'],
            default=['genero', 'raca', 'faixa_etaria', 'escolaridade', 'renda', 'regiao']
        )
        c1, c2 = st.columns(2)
        with c1:
            st.subheader("Conhecimento da PrEP (OR e IC95%)")
            try:
                _, or1 = logistic_or(df, 'Conhecimento_PrEP_bin', features)
                st.dataframe(or1, use_container_width=True)
            except Exception as e:
                st.warning(f"Modelo de conhecimento não ajustado: {e}")
        with c2:
            st.subheader("Uso/propensão a usar PrEP (OR e IC95%)")
            try:
                _, or2 = logistic_or(df, 'Uso_PrEP_bin', features)
                st.dataframe(or2, use_container_width=True)
            except Exception as e:
                st.warning(f"Modelo de uso não ajustado: {e}")

        st.subheader("Perfis latentes (2–5 perfis)")
        try:
            out, best = latent_profiles(df, ['genero','raca','faixa_etaria','escolaridade','regiao','Conhecimento_PrEP_bin','Uso_PrEP_bin'])
            st.write(f"Melhor k (BIC): {best['k']}")
            st.dataframe(out['Perfil_Latente'].value_counts().rename_axis('Perfil').to_frame('N'), use_container_width=True)

            # Tabela resumo por perfil
            with st.expander("Distribuição de variáveis por perfil"):
                tb = pd.crosstab(out['Perfil_Latente'], out['raca'], normalize='index').round(2)
                st.write("Raca x Perfil"); st.dataframe(tb)
        except Exception as e:
            st.warning(f"Perfis não ajustados: {e}")

# ----------------- GAP -----------------
elif menu == "Gap por Grupo":
    st.header("Gap de representação (share PrEP - share HIV) com IC95%")
    df_sp = st.session_state.get('PREP_SP', None)
    df_cases_raw = st.session_state.get('CASES_REF_RAW', None)

    if df_sp is None:
        st.warning("Carregue primeiro o Painel PrEP na aba 'Integração de Dados'.")
    if df_cases_raw is None:
        st.warning("Faça upload do CSV de proxy na aba 'Integração de Dados'.")

    if (df_sp is not None) and (df_cases_raw is not None):
        ano = st.slider("Ano", 2018, 2024, 2024)
        dims_sugeridas = [c for c in ['RACA_COR', 'FAIXA_ETARIA', 'POPULACAO_CHAVE', 'ESCOLARIDADE'] if c in df_sp.columns]
        dim = st.selectbox("Dimensão (coluna em comum)", dims_sugeridas)

        cols_csv = df_cases_raw.columns.tolist()
        col_ano = st.selectbox("Coluna do ANO (CSV)", options=cols_csv, index=0)
        col_dim = st.selectbox("Coluna de DIMENSÃO (CSV)", options=cols_csv, index=min(1, len(cols_csv)-1))
        col_cnt = st.selectbox("Coluna de COUNT (CSV)", options=cols_csv, index=min(2, len(cols_csv)-1))

        df_cases_agg = df_cases_raw.rename(columns={col_ano: 'ANO', col_dim: dim, col_cnt: 'COUNT'}).copy()
        if df_cases_agg[dim].dtype == 'object':
            df_cases_agg[dim] = df_cases_agg[dim].astype(str).str.strip()

        try:
            m = gap_repr_with_ci(df_sp, df_cases_agg, dim_col=dim, ano=ano)
            df_plot = m.copy()
            df_plot['sinal'] = df_plot['gap_repr'].apply(lambda v: 'Sub-representação' if v < 0 else 'Sobre-representação')
            fig = px.bar(df_plot, x='gap_repr', y=dim, orientation='h',
                         color='sinal',
                         error_x='ci_high', error_x_minus='ci_low',
                         color_discrete_map={'Sub-representação': 'crimson', 'Sobre-representação': 'seagreen'},
                         title=f"Gap de representação — {dim} ({ano}) [com IC95%]")
            fig.update_layout(xaxis_title="gap_repr (share_prep - share_hiv)", yaxis_title=dim)
            st.plotly_chart(fig, use_container_width=True)
            st.dataframe(m, use_container_width=True)

            if st.button("Exportar gráfico (PNG)"):
                path = save_fig(fig, f"gap_{dim}_{ano}.png")
                st.success(f"Salvo em {path}")
                download_button_from_file(path, "⬇️ Baixar PNG")

        except Exception as e:
            st.error(f"Erro no cálculo do gap: {e}")

# ----------------- MAPAS -----------------
elif menu == "Mapas":
    st.header("Mapas — Subprefeituras de São Paulo (geopandas)")
    up_geo = st.file_uploader("Opcional: envie um GeoJSON/SHP (zip) oficial do GeoSampa para substituir o padrão", type=['geojson','json','zip','gpkg'])
    gdf = ensure_geodata(default=(up_geo is None), uploaded_file=up_geo)
    st.caption(f"Geometrias carregadas: {len(gdf)} subprefeituras")

    df = harmonize_survey(load_respostas())
    st.markdown("### Choropleth por Subprefeitura")
    if 'subprefeitura' in df.columns and df['subprefeitura'].notna().any():
        agg = (
            df.dropna(subset=['subprefeitura'])
              .groupby('subprefeitura').size().rename('valor').reset_index()
              .rename(columns={'subprefeitura':'Subprefeitura'})
        )
        gviz = choropleth_subpref(gdf, agg, key_df='Subprefeitura', key_gdf='SUBPREF', value_col='valor')
        fig, ax = plt.subplots(1,1,figsize=(8,8))
        gviz.plot(column='valor', legend=True, cmap='YlOrRd', ax=ax, edgecolor='grey')
        ax.set_axis_off()
        st.pyplot(fig)

        if st.button("Exportar mapa (PNG)"):
            fname = f"mapa_subpref_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            path = EXPORT_DIR / fname
            fig.savefig(path, dpi=150, bbox_inches='tight')
            st.success(f"Salvo em {path}")
            download_button_from_file(str(path), "⬇️ Baixar PNG")
    else:
        st.info("Nenhuma coluna 'Subprefeitura' detectada. Informe no formulário ou envie um CSV de mapeamento.")

# ----------------- RELATÓRIO -----------------
elif menu == "Relatório":
    st.header("Relatório .docx (com figuras exportadas)")
    bullets = [
        "Gaps de representação por grupos (share PrEP - share HIV) com IC95%.",
        "Perfis latentes destacam perfis prioritários para comunicação/adesão.",
        "Recomendações: ativar SPrEP e unidades regionais com maior lacuna; ajustar linguagem/canais."
    ]
    salvar = st.text_input("Nome do arquivo (.docx)", value="Relatorio_Prep_SP.docx")
    figs = [str(p) for p in EXPORT_DIR.glob("*.png")]
    if figs:
        st.write("Figuras detectadas para incluir:")
        for f in figs: st.write("•", Path(f).name)
    else:
        st.info("Nenhum PNG encontrado em data/exports. Exporte gráficos nas abas anteriores para incluí-los.")

    if st.button("Gerar relatório .docx"):
        build_report(salvar, bullets, figuras_png=figs)
        st.success(f"Relatório gerado: {salvar}")
        download_button_from_file(salvar, "⬇️ Baixar Relatório DOCX")

# ----------------- SOBRE -----------------
else:
    st.header("Sobre este app")
    st.markdown("""
- **Dados**: Painel PrEP (MS); Painéis de Indicadores (MS); Boletim HIV/aids; PMSP/IST-Aids; GeoSampa.
- **Modelos**: Regressão logística (OR/IC95%); Perfis latentes (GaussianMixture/BIC).
- **Métrica**: Gap de representação (share PrEP - share HIV) com IC95%.

**Observações**  
- Interprete gaps com cautela em amostras pequenas e estratos com baixa contagem.  
- Você pode substituir o GeoJSON padrão por um arquivo oficial do **GeoSampa** (SHP/GeoJSON zip).  
""")